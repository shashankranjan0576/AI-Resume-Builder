from docx import Document


def create_resume_docx(content, filename):

    doc = Document()

    doc.add_heading(
        "ATS Optimized Resume",
        level=1
    )

    doc.add_paragraph(content)

    doc.save(filename)


from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet


def create_resume_pdf(content, filename):

    doc = SimpleDocTemplate(filename)

    styles = getSampleStyleSheet()

    elements = []

    title = Paragraph(
        "ATS Optimized Resume",
        styles['Title']
    )

    elements.append(title)

    elements.append(Spacer(1, 12))

    paragraphs = content.split("\n")

    for para in paragraphs:

        if para.strip():

            p = Paragraph(
                para,
                styles['BodyText']
            )

            elements.append(p)

            elements.append(Spacer(1, 8))

    doc.build(elements)

def clean_ai_output(text):

    unwanted_phrases = [

        "ats-optimized resume",
        "professional cover letter",
        "rewritten version",
        "let me know if you need",
        "i've kept the formatting",
        "i have kept the formatting",
        "tailored to the target role",
        "this resume is optimized",
        "thank you for considering",
        "here's",
        "here is",
        "note:"
    ]

    cleaned_lines = []

    lines = text.split("\n")

    for line in lines:

        should_skip = False

        for phrase in unwanted_phrases:

            if phrase.lower() in line.lower():

                should_skip = True
                break

        if not should_skip:

            line = line.replace("**", "")

            # Convert only + bullets
            if line.strip().startswith("+"):

                line = line.replace("+", "•", 1)

            cleaned_lines.append(line)

    cleaned_text = "\n".join(cleaned_lines)
    cleaned_text = cleaned_text.replace("**", "")
    cleaned_text = cleaned_text.replace("* ", "")
    cleaned_text = cleaned_text.replace("#", "")

    return cleaned_text.strip()